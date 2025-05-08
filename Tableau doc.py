# app.py

import streamlit as st
import xml.etree.ElementTree as ET
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re

# --- Helper Functions ---

def clean_field_name(name):
    """Cleans Tableau's internal field name for display."""
    match = re.match(r"\[(.*?)\]\.\[(.*?)\]", name)
    if match:
        return match.group(2)
    return name.replace("[", "").replace("]", "")

def get_datasource_details(root):
    """Parses all datasources and their columns."""
    datasources_data = {}
    for ds_node in root.findall('.//datasource'):
        ds_name = ds_node.get('name') or ds_node.get('caption')
        if not ds_name: # Fallback for 'federated' datasources without explicit name/caption
            ds_name = ds_node.get('formatted-name', 'Unknown Datasource')

        columns = []
        for col_node in ds_node.findall('.//column'):
            col_name = col_node.get('name') or col_node.get('caption')
            col_role = col_node.get('role')
            col_datatype = col_node.get('datatype')
            col_type = col_node.get('type') # nominal, quantitative, ordinal
            
            calculation_node = col_node.find('calculation')
            formula = calculation_node.get('formula') if calculation_node is not None else None
            
            columns.append({
                'name': clean_field_name(col_name),
                'original_name': col_name, # Keep original for matching
                'role': col_role,
                'datatype': col_datatype,
                'type': col_type,
                'is_calculated': bool(formula),
                'formula': formula
            })
        datasources_data[ds_name] = columns
    return datasources_data

def get_worksheet_details(ws_node, all_datasources_info):
    """Parses details for a single worksheet."""
    ws_name = ws_node.get('name')
    details = {
        'name': ws_name,
        'datasources_used': set(),
        'dimensions': [],
        'measures': [],
        'calculated_fields_used': [],
        'filters': [],
        'fields_on_shelves': [] # For rows, columns, marks
    }

    # Identify datasources used by this worksheet
    ws_datasources = {} # name -> alias or name
    for ds_dep_node in ws_node.findall('.//datasource-dependencies'):
        ds_name = ds_dep_node.get('datasource')
        if ds_name:
            details['datasources_used'].add(ds_name)
            # Find the actual columns for this datasource
            if ds_name in all_datasources_info:
                 ws_datasources[ds_name] = all_datasources_info[ds_name]
            elif ds_dep_node.get('caption') and ds_dep_node.get('caption') in all_datasources_info: # try caption
                ds_name_from_caption = ds_dep_node.get('caption')
                details['datasources_used'].add(ds_name_from_caption)
                ws_datasources[ds_name_from_caption] = all_datasources_info[ds_name_from_caption]


    # Helper to find field details from worksheet's datasources
    def find_field_in_ws_datasources(field_name_to_find):
        cleaned_field_name_to_find = clean_field_name(field_name_to_find)
        for ds_cols in ws_datasources.values():
            for col_data in ds_cols:
                if col_data['name'] == cleaned_field_name_to_find or col_data['original_name'] == field_name_to_find:
                    return col_data
        return None # Not found in this worksheet's specific datasources

    # Parse filters
    for filter_node in ws_node.findall('.//filter'):
        field_name = filter_node.get('column')
        if field_name:
            field_detail = find_field_in_ws_datasources(field_name)
            filter_info = {
                'field': clean_field_name(field_name),
                'class': filter_node.get('class'), # e.g., 'categorical', 'quantitative'
                'datatype': field_detail.get('datatype') if field_detail else 'N/A'
            }
            # Try to get members for categorical filters (simplified)
            members = []
            for member_node in filter_node.findall('.//member'):
                members.append(member_node.get('value'))
            if members:
                filter_info['members'] = members
            details['filters'].append(filter_info)
    
    # Parse shelves (rows, cols, marks card which includes filters there too)
    shelf_types = {
        'rows': 'Rows',
        'cols': 'Columns',
        'color': 'Marks - Color',
        'size': 'Marks - Size',
        'label': 'Marks - Label',
        'detail': 'Marks - Detail',
        'tooltip': 'Marks - Tooltip',
        'shape': 'Marks - Shape',
        'angle': 'Marks - Angle',
        'filter': 'Filters Shelf' # Filters can also appear on marks card shelves
    }

    for column_instance_node in ws_node.findall('.//view/.//column-instance'): # More general way to find fields on shelves
        field_name = column_instance_node.get('column')
        shelf_type_raw = column_instance_node.get('type') # e.g. 'quantitative', 'nominal', 'ordinal'
        # The actual shelf is usually found in parent 'shelf-item' or ancestor 'pane'
        # This is a simplification, getting precise shelf can be tricky.
        # For now, we'll just list fields found in the view section.
        
        parent_shelf = column_instance_node.getparent()
        shelf_name = "Unknown Shelf"
        if parent_shelf is not None and parent_shelf.tag == 'shelf-item':
            shelf_name = parent_shelf.get('name', "Unknown Shelf Item") # e.g. [MarkShelf].[ColorShelf]
            # Try to map to a friendlier name
            for key, friendly_name in shelf_types.items():
                if key in shelf_name.lower():
                    shelf_name = friendly_name
                    break
        elif parent_shelf is not None: # could be 'rows', 'cols' directly under 'pane'
            shelf_name = parent_shelf.tag # 'rows' or 'cols'
            if shelf_name in shelf_types:
                shelf_name = shelf_types[shelf_name]


        if field_name:
            field_detail = find_field_in_ws_datasources(field_name)
            if field_detail:
                shelf_entry = {
                    'field': field_detail['name'],
                    'role': field_detail['role'],
                    'datatype': field_detail['datatype'],
                    'shelf': shelf_name,
                    'type_on_shelf': shelf_type_raw # how it's used (e.g. discrete/continuous)
                }
                details['fields_on_shelves'].append(shelf_entry)

                if field_detail['is_calculated']:
                    # Avoid duplicates if already listed globally
                    if not any(cf['name'] == field_detail['name'] for cf in details['calculated_fields_used']):
                        details['calculated_fields_used'].append(field_detail)
                elif field_detail['role'] == 'dimension':
                    if not any(d['name'] == field_detail['name'] for d in details['dimensions']):
                         details['dimensions'].append(field_detail)
                elif field_detail['role'] == 'measure':
                    if not any(m['name'] == field_detail['name'] for m in details['measures']):
                        details['measures'].append(field_detail)
    
    # Deduplicate (important if a field is on multiple shelves)
    details['dimensions'] = [dict(t) for t in {tuple(d.items()) for d in details['dimensions']}]
    details['measures'] = [dict(t) for t in {tuple(d.items()) for d in details['measures']}]
    details['calculated_fields_used'] = [dict(t) for t in {tuple(d.items()) for d in details['calculated_fields_used']}]


    return details

def parse_twb(xml_content):
    """Main parsing function."""
    root = ET.fromstring(xml_content)
    workbook_docs = []

    all_datasources_info = get_datasource_details(root)
    
    # Create a lookup for all worksheet nodes by name
    all_worksheet_nodes = {ws.get('name'): ws for ws in root.findall('.//worksheet')}

    for db_node in root.findall('.//dashboard'):
        db_name = db_node.get('name')
        dashboard_info = {
            'name': db_name,
            'worksheets': [],
            'objects': [] # For other dashboard objects like text, images (names only)
        }

        # Find worksheets within this dashboard
        # Worksheets are typically inside zones
        for zone_node in db_node.findall('.//zone'):
            worksheet_name_in_zone = zone_node.get('name') # This is the name of the worksheet
            if worksheet_name_in_zone and zone_node.get('type') == 'worksheet':
                if worksheet_name_in_zone in all_worksheet_nodes:
                    ws_node = all_worksheet_nodes[worksheet_name_in_zone]
                    ws_details = get_worksheet_details(ws_node, all_datasources_info)
                    dashboard_info['worksheets'].append(ws_details)
                else:
                    st.warning(f"Worksheet '{worksheet_name_in_zone}' referenced in dashboard '{db_name}' not found in workbook.")
            
            # Capture other object names (text, images, web pages, etc.)
            obj_name = zone_node.get('name')
            obj_type = zone_node.get('type')
            param_name = zone_node.get('param') # For parameters/filters displayed directly
            
            if obj_type and obj_type != 'worksheet':
                name_to_add = obj_name
                if obj_type == 'layout-basic': # Often containers
                    continue # Skip generic containers unless they have specific content
                if not name_to_add and param_name: # e.g. <zone ... type='filter' param='[Parameters].[Parameter 1]' ... />
                    name_to_add = clean_field_name(param_name)
                
                if name_to_add:
                     dashboard_info['objects'].append({'name': name_to_add, 'type': obj_type})


        workbook_docs.append(dashboard_info)
    
    return workbook_docs, all_datasources_info


# --- Output Generation ---

def generate_excel(docs_data):
    """Generates an Excel file from the parsed data."""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        for i, dashboard in enumerate(docs_data):
            df_data = []
            # Sheet name limit is 31 chars, make unique if dashboards have similar long names
            sheet_name = re.sub(r'[\[\]\*:\\\?\/]', '', dashboard['name'])[:25] + f"_{i}"


            df_data.append({'Section': 'Dashboard Info', 'Item': 'Name', 'Details': dashboard['name']})
            if dashboard['objects']:
                 df_data.append({'Section': 'Dashboard Info', 'Item': 'Other Objects', 
                                 'Details': ", ".join([f"{obj['name']} ({obj['type']})" for obj in dashboard['objects']])})


            for ws in dashboard['worksheets']:
                ws_header = f"Worksheet: {ws['name']}"
                df_data.append({'Section': ws_header, 'Item': 'Datasources', 'Details': ", ".join(list(ws['datasources_used']))})

                for dim in ws['dimensions']:
                    df_data.append({'Section': ws_header, 'Item': 'Dimension', 'Details': f"{dim['name']} (Type: {dim['datatype']})"})
                for meas in ws['measures']:
                    df_data.append({'Section': ws_header, 'Item': 'Measure', 'Details': f"{meas['name']} (Type: {meas['datatype']})"})
                for cf in ws['calculated_fields_used']:
                    df_data.append({'Section': ws_header, 'Item': 'Calculated Field', 'Details': f"{cf['name']} (Formula: {cf['formula']})"})
                for filt in ws['filters']:
                    members_str = f" (Members: {', '.join(filt.get('members', []))})" if filt.get('members') else ""
                    df_data.append({'Section': ws_header, 'Item': 'Filter', 'Details': f"{filt['field']} (Type: {filt.get('class', 'N/A')}){members_str}"})
                for shelf_item in ws['fields_on_shelves']:
                     df_data.append({'Section': ws_header, 'Item': f"Field on Shelf ({shelf_item['shelf']})", 
                                     'Details': f"{shelf_item['field']} (Role: {shelf_item['role']}, Datatype: {shelf_item['datatype']}, Usage: {shelf_item['type_on_shelf']})"})
            
            df = pd.DataFrame(df_data)
            if not df.empty:
                df.to_excel(writer, sheet_name=sheet_name, index=False)
    return output.getvalue()

def generate_word(docs_data):
    """Generates a Word document from the parsed data."""
    doc = Document()
    doc.add_heading('Tableau Workbook Documentation', level=0)

    for dashboard in docs_data:
        doc.add_heading(f"Dashboard: {dashboard['name']}", level=1)
        
        # Placeholder for snippet instructions
        p = doc.add_paragraph()
        p.add_run("Instructions for Dashboard Snippet:").bold = True
        doc.add_paragraph(
            "1. Take a screenshot of this dashboard in Tableau.\n"
            "2. Paste the screenshot here.\n"
            "3. Use the component list below to manually number sections on your screenshot."
        )
        doc.add_paragraph(f"Dashboard Components (for manual numbering on snippet):")
        num = 1
        if dashboard['objects']:
            doc.add_paragraph(f"  General Objects:")
            for obj in dashboard['objects']:
                 doc.add_paragraph(f"    {num}. {obj['name']} ({obj['type']})", style='ListNumber')
                 num +=1

        for i_ws, ws in enumerate(dashboard['worksheets']):
            doc.add_paragraph(f"  {num}. Worksheet: {ws['name']}", style='ListNumber')
            num +=1
        doc.add_paragraph("--- End of Component List ---")


        if dashboard['objects']:
            doc.add_heading("Dashboard-Level Objects", level=2)
            for obj in dashboard['objects']:
                doc.add_paragraph(f"- {obj['name']} (Type: {obj['type']})", style='ListBullet')

        for ws in dashboard['worksheets']:
            doc.add_heading(f"Worksheet: {ws['name']}", level=2)
            
            if ws['datasources_used']:
                doc.add_paragraph(f"Datasources: {', '.join(list(ws['datasources_used']))}")

            if ws['dimensions']:
                doc.add_heading("Dimensions Used:", level=3)
                for item in ws['dimensions']:
                    doc.add_paragraph(f"- {item['name']} (Datatype: {item['datatype']})", style='ListBullet')
            
            if ws['measures']:
                doc.add_heading("Measures Used:", level=3)
                for item in ws['measures']:
                    doc.add_paragraph(f"- {item['name']} (Datatype: {item['datatype']})", style='ListBullet')

            if ws['calculated_fields_used']:
                doc.add_heading("Calculated Fields Used:", level=3)
                for item in ws['calculated_fields_used']:
                    doc.add_paragraph(f"{item['name']}", style='ListBullet')
                    p_formula = doc.add_paragraph(f"  Formula: {item['formula']}")
                    p_formula.paragraph_format.left_indent = Inches(0.5)
            
            if ws['filters']:
                doc.add_heading("Filters:", level=3)
                for item in ws['filters']:
                    members_str = f" (Selected: {', '.join(item.get('members', []))})" if item.get('members') else ""
                    doc.add_paragraph(f"- {item['field']} (Type: {item.get('class', 'N/A')}){members_str}", style='ListBullet')

            if ws['fields_on_shelves']:
                doc.add_heading("Fields on Shelves:", level=3)
                for item in ws['fields_on_shelves']:
                     doc.add_paragraph(f"- {item['field']} (On: {item['shelf']}, Role: {item['role']}, Usage: {item['type_on_shelf']})", style='ListBullet')
            
            doc.add_paragraph() # Add some space

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# --- Streamlit App UI ---
st.set_page_config(layout="wide")
st.title("Tableau Workbook Documenter 📄")

st.markdown("""
Upload a Tableau Workbook XML file (`.twb`) to generate documentation.
This tool will extract information about dashboards, worksheets, fields, calculations, and filters.

**Note on Dashboard Snippets:** This tool **cannot** automatically generate screenshots of your dashboards.
It will provide a list of components for each dashboard. You'll need to:
1.  Manually take a screenshot of your dashboard.
2.  Refer to the component list provided in the documentation (especially the Word output) to annotate your screenshot.
""")

uploaded_file = st.file_uploader("Choose a .twb file", type="twb")

if uploaded_file is not None:
    try:
        xml_content = uploaded_file.read()
        st.success(f"File '{uploaded_file.name}' uploaded successfully!")

        with st.spinner("Parsing workbook... this may take a moment for large files."):
            parsed_docs, all_datasources_info = parse_twb(xml_content) # Make sure parse_twb returns this

        if not parsed_docs:
            st.warning("No dashboards found in the workbook or an error occurred during parsing.")
        else:
            st.header("Workbook Documentation Overview")

            # Display Datasource Information
            if all_datasources_info:
                st.subheader("Global Datasource Information")
                for ds_name, ds_cols in all_datasources_info.items():
                    with st.expander(f"Datasource: {ds_name} ({len(ds_cols)} columns)"):
                        # Display only a few columns initially or provide a table
                        df_cols = pd.DataFrame(ds_cols)
                        st.dataframe(df_cols[['name', 'role', 'datatype', 'is_calculated', 'formula']].head(10), height=300)
                        if len(df_cols) > 10:
                            st.caption(f"...and {len(df_cols)-10} more columns.")


            # Display Dashboard by Dashboard
            for i, dashboard_data in enumerate(parsed_docs):
                st.subheader(f"Dashboard {i+1}: {dashboard_data['name']}")

                # Placeholder for Snippet Section
                st.markdown(f"**Visual Snippet Area (Manual)**")
                st.info(f"""
                Please take a screenshot of the '{dashboard_data['name']}' dashboard.
                The downloadable Word document will provide a list of components (see below)
                that you can use to number sections on your screenshot.
                """)
                
                # List components for numbering
                st.markdown("**Dashboard Components (for numbering):**")
                comp_num = 1
                if dashboard_data['objects']:
                    st.markdown("  *General Objects:*")
                    for obj in dashboard_data['objects']:
                        st.markdown(f"    {comp_num}. {obj['name']} ({obj['type']})")
                        comp_num +=1
                
                for ws_idx, ws_data in enumerate(dashboard_data['worksheets']):
                    st.markdown(f"  {comp_num}. *Worksheet:* {ws_data['name']}")
                    comp_num += 1


                # Dashboard-level objects
                if dashboard_data['objects']:
                    with st.expander(f"Dashboard-Level Objects/Controls ({len(dashboard_data['objects'])} items)"):
                        for obj in dashboard_data['objects']:
                            st.write(f"- **{obj['name']}** (Type: {obj['type']})")
                
                # Worksheets in Dashboard
                for ws_idx, ws_data in enumerate(dashboard_data['worksheets']):
                    with st.expander(f"Worksheet: {ws_data['name']}"):
                        st.markdown(f"**Datasources:** {', '.join(list(ws_data['datasources_used'])) if ws_data['datasources_used'] else 'N/A'}")

                        if ws_data['dimensions']:
                            st.markdown("**Dimensions Used:**")
                            for item in ws_data['dimensions']: st.markdown(f"- `{item['name']}` (Type: {item['datatype']})")
                        
                        if ws_data['measures']:
                            st.markdown("**Measures Used:**")
                            for item in ws_data['measures']: st.markdown(f"- `{item['name']}` (Type: {item['datatype']})")

                        if ws_data['calculated_fields_used']:
                            st.markdown("**Calculated Fields Used:**")
                            for item in ws_data['calculated_fields_used']:
                                st.markdown(f"- `{item['name']}`")
                                st.code(f"Formula: {item['formula']}", language='sql') # or 'plaintext'
                        
                        if ws_data['filters']:
                            st.markdown("**Filters:**")
                            for item in ws_data['filters']:
                                members_str = f" (Selected: {', '.join(item.get('members', []))})" if item.get('members') else ""
                                st.markdown(f"- `{item['field']}` (Type: {item.get('class', 'N/A')}){members_str}")
                        
                        if ws_data['fields_on_shelves']:
                            st.markdown("**Fields on Shelves (Rows, Columns, Marks, etc.):**")
                            # Create a small dataframe for better display
                            shelf_df_data = []
                            for item in ws_data['fields_on_shelves']:
                                shelf_df_data.append({
                                    'Field': item['field'], 
                                    'Shelf': item['shelf'], 
                                    'Role': item['role'], 
                                    'Datatype': item['datatype'],
                                    'Usage (Type on Shelf)': item['type_on_shelf']
                                })
                            if shelf_df_data:
                                st.dataframe(pd.DataFrame(shelf_df_data))
                st.markdown("---")

            # Download Buttons
            st.header("Download Documentation")
            col1, col2 = st.columns(2)
            
            with col1:
                excel_data = generate_excel(parsed_docs)
                st.download_button(
                    label="📥 Download as Excel",
                    data=excel_data,
                    file_name=f"{uploaded_file.name.replace('.twb', '')}_documentation.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            with col2:
                word_data = generate_word(parsed_docs)
                st.download_button(
                    label="📄 Download as Word",
                    data=word_data,
                    file_name=f"{uploaded_file.name.replace('.twb', '')}_documentation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    except ET.ParseError:
        st.error("Error parsing the XML file. Please ensure it's a valid .twb file and not corrupted.")
    except Exception as e:
        st.error(f"An unexpected error occurred: {e}")
        st.error("This could be due to a very complex or unusually structured TWB file.")
        import traceback
        st.text(traceback.format_exc()) # For debugging
